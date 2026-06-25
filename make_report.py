#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
기계설비 성능점검 검토자문 시스템 — A4 5페이지 보고서 생성 (DOCX)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── 색상 / 폰트 ───────────────────────────────────────
FONT      = '맑은 고딕'
C_NAVY    = RGBColor(0x15, 0x40, 0x63)
C_GREEN   = RGBColor(0x73, 0xBF, 0x45)
C_RED     = RGBColor(0xEF, 0x41, 0x29)
C_MUTED   = RGBColor(0x55, 0x55, 0x55)
C_GREY    = RGBColor(0x88, 0x88, 0x88)
C_BG_HEAD = 'D6E4F0'   # 표 헤더 배경(hex)
C_BG_ALT  = 'EAF3FB'   # 표 짝수행 배경
C_PLACEHOLDER = 'EBEBEB'

doc = Document()

# ── A4 페이지 설정 ────────────────────────────────────
sec = doc.sections[0]
sec.page_height    = Cm(29.7)
sec.page_width     = Cm(21.0)
sec.left_margin    = Cm(2.5)
sec.right_margin   = Cm(2.5)
sec.top_margin     = Cm(2.5)
sec.bottom_margin  = Cm(2.5)

# ── 스타일 헬퍼 ──────────────────────────────────────

def _rfonts(run, name):
    """한글 폰트를 동시에 설정"""
    try:
        rpr = run._element.get_or_add_rPr()
        rf  = OxmlElement('w:rFonts')
        for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            rf.set(qn(attr), name)
        old = rpr.find(qn('w:rFonts'))
        if old is not None:
            rpr.remove(old)
        rpr.insert(0, rf)
    except Exception:
        pass


def rfmt(run, size=11, bold=False, italic=False, color=None):
    run.font.name   = FONT
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    _rfonts(run, FONT)


def pspacing(para, before=0, after=80, line=360, rule='auto'):
    """단락 간격·줄간격"""
    ppr = para._p.get_or_add_pPr()
    old = ppr.find(qn('w:spacing'))
    if old is not None:
        ppr.remove(old)
    spc = OxmlElement('w:spacing')
    spc.set(qn('w:before'),   str(before))
    spc.set(qn('w:after'),    str(after))
    spc.set(qn('w:line'),     str(line))
    spc.set(qn('w:lineRule'), rule)
    ppr.append(spc)


def cell_shd(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)


def cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    brd  = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:color'), 'AAAAAA')
        brd.append(b)
    tcPr.append(brd)


# ── 단락 추가 헬퍼 ────────────────────────────────────

def add_heading(text, level=1):
    p = doc.add_paragraph()
    pspacing(p, before=200 if level == 1 else 140,
                after=80,
                line=320, rule='auto')
    r = p.add_run(text)
    if level == 1:
        rfmt(r, size=15, bold=True, color=C_NAVY)
    elif level == 2:
        rfmt(r, size=13, bold=True, color=C_NAVY)
    else:
        rfmt(r, size=11, bold=True, color=C_MUTED)
    return p


def add_body(text, indent=0.0, before=0, after=80):
    p = doc.add_paragraph()
    pspacing(p, before=before, after=after, line=360)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    rfmt(r, size=10.5)
    return p


def add_bullet(text, num=None, sub=False):
    p = doc.add_paragraph()
    pspacing(p, before=0, after=50, line=330)
    indent = Cm(1.2) if not sub else Cm(2.0)
    p.paragraph_format.left_indent       = indent
    p.paragraph_format.first_line_indent = Cm(-0.55)
    if num:
        r1 = p.add_run(f'{num}.  ')
        rfmt(r1, size=10.5, bold=True, color=C_NAVY)
    else:
        r1 = p.add_run('•  ' if not sub else '–  ')
        rfmt(r1, size=10.5, bold=False, color=C_GREEN if not sub else C_MUTED)
    r2 = p.add_run(text)
    rfmt(r2, size=10.5)
    return p


def add_placeholder(label, height_cm=6.0):
    """회색 이미지 자리 표시자"""
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)

    # 행 높이
    tr   = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'),   str(int(height_cm * 567)))
    trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)

    # 셀 스타일
    cell_shd(cell, C_PLACEHOLDER)
    cell_borders(cell)

    # 텍스트 (세로 중앙)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(f'\n[ 스크린샷 삽입: {label} ]')
    rfmt(cr, size=11, bold=True, color=C_GREY)

    # 캡션
    gap = doc.add_paragraph()
    pspacing(gap, before=40, after=0, line=240)
    cap = doc.add_paragraph()
    pspacing(cap, before=0, after=120, line=260)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr2 = cap.add_run(f'[그림] {label}')
    rfmt(cr2, size=9.5, italic=True, color=C_MUTED)
    return tbl


def add_gap(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        pspacing(p, before=0, after=0, line=240)


# ════════════════════════════════════════════════════════
# ① 표지 (Page 1)
# ════════════════════════════════════════════════════════
add_gap(3)

p = doc.add_paragraph()
pspacing(p, before=0, after=80, line=300)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('대한기계설비산업연구원 (KRIMFI)')
rfmt(r, size=13, color=C_MUTED)

add_gap(1)

p = doc.add_paragraph()
pspacing(p, before=0, after=120, line=420)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('기계설비 성능점검\n검토자문 시스템 개발 보고서')
rfmt(r, size=24, bold=True, color=C_NAVY)

p = doc.add_paragraph()
pspacing(p, before=0, after=0, line=300)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AI 기반 성능점검 보고서 적정성 자동 검토 플랫폼')
rfmt(r, size=13, italic=True, color=C_GREEN)

add_gap(4)

# 정보 표
info_tbl = doc.add_table(rows=4, cols=2)
info_data = [
    ('플랫폼 명칭', '기계설비 성능점검 검토자문 시스템 (MPIS Reviewer)'),
    ('활용 AI 모델', 'Google Gemini 멀티모달 대형언어모델'),
    ('작 성 기 관', '대한기계설비산업연구원 (KRIMFI)'),
    ('작 성 일 자', datetime.date.today().strftime('%Y년 %m월 %d일')),
]
col_widths = [Cm(4), Cm(11.5)]
for i, (label, value) in enumerate(info_data):
    row = info_tbl.rows[i]
    row.cells[0].width = col_widths[0]
    row.cells[1].width = col_widths[1]

    lc = info_tbl.cell(i, 0)
    vc = info_tbl.cell(i, 1)
    cell_shd(lc, C_BG_HEAD)
    cell_shd(vc, 'FFFFFF')
    cell_borders(lc)
    cell_borders(vc)

    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(label)
    rfmt(lr, size=11, bold=True, color=C_NAVY)

    vp = vc.paragraphs[0]
    vr = vp.add_run(value)
    rfmt(vr, size=11)

doc.add_page_break()


# ════════════════════════════════════════════════════════
# ② 1장. 서론 및 필요성  (Page 2)
# ════════════════════════════════════════════════════════
add_heading('1장. 서론 및 필요성', level=1)

add_heading('1.1 연구 배경', level=2)
add_body(
    '「기계설비법」 제17조에 따라 일정 규모 이상의 건축물은 의무적으로 기계설비 성능점검을 실시하고, '
    '그 결과를 보고서로 제출하여야 한다. 성능점검은 공기조화기, 냉·난방기, 환기설비, 위생설비 등 '
    '건물 내 주요 기계설비를 대상으로 하며, 점검 항목은 장비 종류에 따라 수십~수백 개에 달한다.'
)
add_body(
    '점검업체가 제출한 보고서는 관할 기관 또는 발주처에서 적정성 검토를 수행해야 하나, '
    '현행 검토 방식은 전문 인력이 수작업으로 수치와 사진을 대조하는 방식에 의존하고 있어 '
    '시간적·인적 자원의 낭비가 크다.'
)

add_heading('1.2 현황 및 문제점', level=2)
add_bullet('항목 과다 및 검토 복잡성: 점검 항목이 많고 계측기 사진·수치·판정 결과가 혼재하여 수작업 검토의 부담이 극대화됨')
add_bullet('검토 일관성 부재: 검토 담당자에 따라 판정 기준이 달라져 동일한 결과물에 대해 판정이 상이하게 나타남')
add_bullet('정성적 표현 의존: 점검업체 보고서가 "양호", "적합" 등 정성적 서술에 그쳐 근거 기반 검증이 어려움')
add_bullet('설계값 비교 누락: 측정값과 설계 사양값(허용범위) 간의 정량적 비교가 빠진 항목이 다수 발견됨')

add_gap(1)
add_placeholder('기계설비 성능점검 보고서 예시 화면', height_cm=5.5)


# ════════════════════════════════════════════════════════
# ③ 2장. 시스템 개요 및 목적  (Page 2 → 3)
# ════════════════════════════════════════════════════════
add_heading('2장. 시스템 개요 및 목적', level=1)

add_heading('2.1 시스템 개요', level=2)
add_body(
    '본 시스템은 구글 Gemini 멀티모달 대형언어모델(LLM)을 활용하여 기계설비 성능점검 보고서(PDF 또는 '
    '현장 사진)를 자동으로 분석하고, 항목별 적정성을 판정하는 웹 기반 플랫폼이다. '
    'Python Streamlit 프레임워크로 구현되어 별도 설치 없이 웹 브라우저에서 접근 가능하다.'
)

add_heading('2.2 개발 목적', level=2)
add_bullet('점검결과의 적정성을 AI가 자동 검토하여 [적정] / [보완필요]로 판정함으로써 수작업 검토를 대체')
add_bullet('설계값·측정값·허용 범위를 포함한 정량적·기술적 자문의견을 자동 생성하여 근거 중심 검토 실현')
add_bullet('"양호·적합" 등 정성 표현에 머무르지 않는 수치 기반 자문의견 제공으로 검토 품질 균질화')
add_bullet('검토 결과를 화면 분류 및 엑셀 파일로 자동 구조화하여 행정 효율성 향상')
add_bullet('검토 데이터의 디지털 축적을 통해 건물 기계설비 성능 이력 관리 기반 마련')


# ════════════════════════════════════════════════════════
# ④ 3장. 시스템 프로세스  (Page 3 → 4)
# ════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('3장. 시스템 프로세스', level=1)

add_body(
    '본 플랫폼의 핵심 프로세스는 [그림 3-1]과 같이 ① 보고서 업로드 → ② AI 멀티모달 분석 → '
    '③ 검토원칙 적용 → ④ 결과 구조화의 4단계로 구성된다.'
)

add_placeholder('시스템 메인 화면 전체 (MPIS Reviewer 대시보드)', height_cm=6.0)

add_heading('3.1 1단계: 보고서 업로드', level=2)
add_body(
    '사용자는 좌측 사이드바를 통해 점검업체가 제출한 보고서를 업로드한다. '
    '지원 형식은 PDF 문서와 JPG/PNG 사진 파일이며, 복수의 파일을 동시에 업로드할 수 있어 '
    '현장에서 촬영한 계측기 사진을 다중 첨부하는 것이 가능하다.'
)
add_bullet('지원 형식: PDF 문서(복수 페이지), JPG / PNG 계측기 사진(다중 선택)')
add_bullet('사용 AI 모델: Google Gemini API를 통해 Pro / Flash 등 사용 가능 모델 중 선택')

add_placeholder('좌측 사이드바 — 파일 업로드 및 모델 선택 화면', height_cm=5.5)

add_heading('3.2 2단계: AI 멀티모달 분석', level=2)
add_body(
    '"검토자문서 생성" 버튼을 클릭하면 선택된 Gemini 모델이 업로드된 문서 전체를 분석한다. '
    '텍스트와 이미지를 동시에 처리하는 멀티모달 능력을 활용하여 계측기 사진 속 수치를 직접 판독하고, '
    '보고서에 기재된 모든 장비와 점검 항목을 누락 없이 추출한다.'
)
add_bullet('전 장비·전 점검항목 100% 추출 (누락 금지 시스템 프롬프트 적용)')
add_bullet('계측기 사진 내 수치 자동 판독 (멀티모달 이미지 분석)')
add_bullet('불완전한 AI 응답 자동 복구 로직 내장 (JSON 파싱 오류 대응)')


# ════════════════════════════════════════════════════════
# ⑤ 3장 계속 + 4장  (Page 4 → 5)
# ════════════════════════════════════════════════════════
add_heading('3.3 3단계: 검토원칙 적용 및 판정', level=2)
add_body(
    'AI는 추출된 각 점검 항목에 대해 아래 5가지 정량 검토 원칙을 적용하여 '
    '[적정] 또는 [보완필요]로 판정하고, 근거를 서술한 전문가 검토의견을 생성한다.'
)

# 5원칙 표
prin_tbl = doc.add_table(rows=6, cols=2)
prin_data = [
    ('원칙', '내용'),
    ('원칙 1', '설계값 대비 측정값 비율 계산·명시  (예: 측정 18,634 CMH = 설계값의 88.73%)'),
    ('원칙 2', '단위 환산 직접 수행  (풍속 m/s → 풍량 CMH, 단면적 활용 계산)'),
    ('원칙 3', '운전조건 확인  (정격운전 / 일반운전 구분 후 기준값 적용)'),
    ('원칙 4', '판정 근거를 정량 논리로 서술  (정성 표현 사용 금지)'),
    ('원칙 5', '보완 필요 시 구체적 개선방향 제시  (조치 항목 명시)'),
]
for i, (k, v) in enumerate(prin_data):
    lc = prin_tbl.cell(i, 0)
    vc = prin_tbl.cell(i, 1)
    cell_borders(lc)
    cell_borders(vc)
    if i == 0:
        cell_shd(lc, C_BG_HEAD)
        cell_shd(vc, C_BG_HEAD)
    elif i % 2 == 0:
        cell_shd(lc, C_BG_ALT)
        cell_shd(vc, C_BG_ALT)
    lp = lc.paragraphs[0]; lr = lp.add_run(k)
    rfmt(lr, size=10.5, bold=(i == 0), color=C_NAVY if i == 0 else None)
    vp = vc.paragraphs[0]; vr = vp.add_run(v)
    rfmt(vr, size=10.5, bold=(i == 0), color=C_NAVY if i == 0 else None)

add_gap(1)

add_heading('3.4 4단계: 결과 구조화 및 출력', level=2)
add_body(
    '판정 결과는 화면에 [적정] / [보완필요] 탭으로 분류되어 즉시 표시되며, '
    '"엑셀 다운로드" 기능을 통해 전체 검토 결과를 구조화된 스프레드시트로 내보낼 수 있다. '
    '또한 PPTX 보고서 자동 생성 기능을 통해 프레젠테이션용 결과물을 산출한다.'
)
add_bullet('[적정] / [보완필요] 탭 분류 및 항목별 상세 의견 표시')
add_bullet('전체 검토 결과 엑셀(.xlsx) 다운로드')
add_bullet('결과 요약 PPTX 자동 생성')

add_placeholder('검토 결과 화면 — [적정] / [보완필요] 탭 분류', height_cm=5.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# ⑥ 4장. 실증 결과  (Page 5)
# ════════════════════════════════════════════════════════
add_heading('4장. 실증 결과 및 기대효과', level=1)

add_heading('4.1 충무아트센터 성능점검 보고서 실증', level=2)
add_body(
    '본 시스템의 성능 검증을 위해 서울 충무아트센터의 기계설비 성능점검 보고서를 '
    '대상으로 실증을 수행하였다. 대상 장비는 공기조화기(AHU-5, 10, 15, 17), '
    '히트펌프유닛(HVU-1), 팬코일유닛 등이며, 총 123개 점검 항목에 대해 검토를 실시하였다.'
)

# 실증 결과 표
res_tbl = doc.add_table(rows=3, cols=3)
res_data = [
    ('구분', '건수', '비율'),
    ('적정', '64건', '52%'),
    ('보완필요', '59건', '48%'),
]
for i, row_data in enumerate(res_data):
    for j, val in enumerate(row_data):
        cell = res_tbl.cell(i, j)
        cell_borders(cell)
        if i == 0:
            cell_shd(cell, C_BG_HEAD)
        elif i == 2:
            cell_shd(cell, 'FFF0EE')
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(val)
        bold_flag = (i == 0) or (j == 1 and i > 0)
        color_flag = C_NAVY if i == 0 else (C_RED if i == 2 and j == 1 else None)
        rfmt(cr, size=11, bold=bold_flag, color=color_flag)

add_gap(1)

add_heading('4.2 대표 사례: AHU-5 송풍기 풍량 판정 정정', level=2)
add_body(
    '점검업체는 AHU-5 송풍기 풍량을 별도의 정량 비교 없이 "적합"으로 판정하였다. '
    '그러나 AI 시스템은 측정값(18,634 CMH)이 설계 풍량(21,000 CMH)의 88.73%에 불과하여 '
    '허용 기준(±10%, 90~110%)을 충족하지 못함을 계산하고 [보완필요]로 재판정하였다. '
    '이는 본 시스템이 단순 결과 수용이 아닌 정량 근거 기반 검토를 수행함을 보여 준다.'
)

add_placeholder('검토 결과 상세 화면 — AHU-5 항목 보완필요 판정 예시', height_cm=5.0)

add_heading('4.3 기대효과', level=2)
add_bullet('검토 일관성: 표준화된 AI 기준 적용으로 검토자 간 판정 편차 제거')
add_bullet('업무 효율: 수십~수백 개 항목의 자동 검토로 검토 시간 90% 이상 단축 가능')
add_bullet('데이터 축적: 디지털 검토 이력 저장을 통한 건물 성능 데이터베이스 구축')
add_bullet('품질 향상: 설계값 대비 측정값 정량 비교 의무화로 보고서 품질 수준 향상 유도')

add_gap(1)

# 결론
add_heading('결론', level=1)
add_body(
    '본 보고서에서는 기계설비 성능점검 보고서의 적정성을 AI가 자동으로 검토·판정하는 '
    '"기계설비 성능점검 검토자문 시스템"의 개발 배경, 목적, 프로세스, 실증 결과를 기술하였다. '
    '충무아트센터 실증을 통해 시스템의 정량 검토 능력 및 판정 신뢰성이 확인되었으며, '
    '향후 다양한 건물 유형과 점검 보고서에 확장 적용하여 기계설비 성능점검 검토의 표준화에 '
    '기여할 것으로 기대된다.'
)

# ── 저장 ─────────────────────────────────────────────
OUT = (r'c:\Users\이창재\Desktop\py코딩연습\mpis-reviewer'
       r'\기계설비_성능점검_검토자문시스템_보고서.docx')
doc.save(OUT)
print(f'저장 완료: {OUT}')
